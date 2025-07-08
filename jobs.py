import json
import os
import shutil
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

def create_resume(json_path, company, position):
    with open(json_path, 'r') as file:
        data = json.load(file)

    # Output folder
    output_folder = "resumes"
    os.makedirs(output_folder, exist_ok=True)

    # Temp folder to hold docx
    temp_folder = "temp_docx"
    os.makedirs(temp_folder, exist_ok=True)

    # Create DOCX
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15

    # Name (bold, large, centered)
    name_para = doc.add_paragraph()
    run = name_para.add_run(data["name"])
    run.bold = True
    run.font.size = Pt(16)
    name_para.paragraph_format.space_after = Pt(2)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact Info (centered, single line)
    contact_items = [
        data["contact"]["email"],
        data["contact"]["phone"],
        data["contact"]["location"]
    ]
    if data["contact"].get("linkedin"):
        contact_items.append(data["contact"]["linkedin"])
    contact_info = " | ".join(contact_items)
    contact_para = doc.add_paragraph(contact_info)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # EDUCATION
    edu_heading = doc.add_paragraph()
    edu_run = edu_heading.add_run("EDUCATION")
    edu_run.bold = True
    edu_heading.paragraph_format.space_before = Pt(0)
    edu_heading.paragraph_format.space_after = Pt(8)  # <--- Only change made here

    for i, edu in enumerate(data["education"]):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        if i == len(data["education"]) - 1:
            p.paragraph_format.space_after = Pt(8)
        else:
            p.paragraph_format.space_after = Pt(0)
        line = f"{edu['degree']}, {edu['institution']} | {edu['graduation_year']}"
        p.add_run(line).bold = True

    # EXPERIENCE
    exp_heading = doc.add_paragraph()
    exp_run = exp_heading.add_run("EXPERIENCE")
    exp_run.bold = True
    exp_heading.paragraph_format.space_before = Pt(0)
    exp_heading.paragraph_format.space_after = Pt(0)

    for job_index, job in enumerate(data["experience"]):
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(0)
        left = f"{job['position']} | {job['company']}"
        right = f"{job['start_date']} - {job['end_date']}"
        p.add_run(left).bold = True
        p.add_run("\t")
        p.add_run(right).bold = True

        for bullet_index, bullet in enumerate(job['bullets']):
            bullet_para = doc.add_paragraph(bullet, style='List Bullet')
            bullet_para.paragraph_format.left_indent = Inches(0.25)
            bullet_para.paragraph_format.space_before = Pt(0)

            if job_index == len(data["experience"]) - 1 and bullet_index == len(job['bullets']) - 1:
                bullet_para.paragraph_format.space_after = Pt(8)
            else:
                bullet_para.paragraph_format.space_after = Pt(0)

    # SUMMARY
    summary_heading = doc.add_paragraph()
    summary_run = summary_heading.add_run("SUMMARY")
    summary_run.bold = True
    summary_heading.paragraph_format.space_before = Pt(0)
    summary_heading.paragraph_format.space_after = Pt(0)
    doc.add_paragraph(data["summary"])

    # SKILLS
    skills_heading = doc.add_paragraph()
    skills_run = skills_heading.add_run("SKILLS")
    skills_run.bold = True
    skills_heading.paragraph_format.space_before = Pt(0)
    skills_heading.paragraph_format.space_after = Pt(0)

    if isinstance(data["skills"], dict):
        all_skills = []
        for cat, skills in data["skills"].items():
            all_skills.extend(skills)
        skills_text = ", ".join(all_skills)
    else:
        skills_text = ", ".join(data["skills"])
    doc.add_paragraph(skills_text)

    # PROJECTS
    if "projects" in data and data["projects"]:
        proj_heading = doc.add_paragraph()
        proj_run = proj_heading.add_run("PROJECTS")
        proj_run.bold = True
        proj_heading.paragraph_format.space_before = Pt(0)
        proj_heading.paragraph_format.space_after = Pt(0)

        for project in data["projects"]:
            title = doc.add_paragraph(project["title"])
            title.paragraph_format.space_before = Pt(0)
            title.paragraph_format.space_after = Pt(0)
            title.runs[0].bold = True
            desc = doc.add_paragraph(project["description"])
            desc.paragraph_format.space_before = Pt(0)
            desc.paragraph_format.space_after = Pt(0)
            if "technologies" in project:
                tech = doc.add_paragraph("Technologies: " + ", ".join(project["technologies"]))
                tech.paragraph_format.space_before = Pt(0)
                tech.paragraph_format.space_after = Pt(0)

    # Save DOCX temporarily
    temp_docx_path = os.path.join(temp_folder, "resume.docx")
    doc.save(temp_docx_path)

    # Convert DOCX to PDF in output folder
    pdf_output_path = os.path.join(output_folder, f"{company}_{position}_resume.pdf")
    try:
        convert(temp_docx_path, pdf_output_path)
        print(f"✅ PDF saved to {pdf_output_path}")
    except Exception as e:
        print(f"[!] PDF conversion failed: {e}")

    # Cleanup temp files
    shutil.rmtree(temp_folder)

# Example usage
if __name__ == "__main__":
    create_resume("darius_resume.json", "tailored", "")
